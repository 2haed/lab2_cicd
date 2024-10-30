import pytest
from app.main import process_file, extract_images_from_docx
from docx import Document
import docx
from PIL import Image
import io
import os

@pytest.fixture
def docx_file_with_image(tmp_path):
    file = tmp_path / "test_file_with_image.docx"
    doc = Document()
    doc.add_paragraph("This is a sample docx file with text and one image.")

    image = Image.new('RGB', (100, 100), color='red')
    image_stream = io.BytesIO()
    image.save(image_stream, format='JPEG')
    image_stream.seek(0)
    
    doc.add_picture(image_stream, width=docx.shared.Inches(1))
    doc.save(file)
    return file

def test_process_and_extract_images(docx_file_with_image):
    text_content = process_file(str(docx_file_with_image))
    assert "This is a sample docx file with text and one image." in text_content

    output_folder = 'app/data'
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    image_count = extract_images_from_docx(str(docx_file_with_image))

    assert image_count == 1
    assert os.path.exists(os.path.join(output_folder, "image_1.jpeg"))

    os.remove(os.path.join(output_folder, "image_1.jpeg"))