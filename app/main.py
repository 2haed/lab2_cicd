from docx import Document
from PIL import Image
import io

def process_file(file_path):
    if file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Не поддерживаемый формат")
    return content

def count_words(text):
    return len(text.split())

def extract_images_from_docx(file_path):
    doc = Document(file_path)
    image_count = 0

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            image_data = rel.target_part.blob
            image = Image.open(io.BytesIO(image_data))
            image_format = image.format.lower()
            image.save(f"app/data/image_{image_count}.{image_format}")

    return image_count

extract_images_from_docx('test.docx')